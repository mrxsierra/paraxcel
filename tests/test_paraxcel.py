import pytest
from pydantic import ValidationError
from docx import Document as docx_instance
from docx.document import Document
from docx.enum.text import WD_COLOR
from src.paraxcel.docx_parser import parse_para, read_docx
from src.paraxcel.model import Question
from src.paraxcel.para_utility import inject_html_font_formats, remove_prefix, find_marked_answer
from src.paraxcel.excel_writer import para_to_excel
import pandas as pd
import os

# Test the Question model validation
def test_question_model_valid():
    question = Question(
        question_text="What is the capital of France?",
        options=["Berlin", "Madrid", "Paris", "Rome"],
        correct_answer=2
    )
    assert question.correct_answer == 2

def test_question_model_invalid():
    with pytest.raises(ValidationError):
        Question(
            question_text="What is the capital of France?",
            options=["Berlin", "Madrid", "Paris", "Rome"],
            correct_answer=5  # Out of range (valid is 0-3)
        )

def test_question_model_none_correct_answer():
    question = Question(
        question_text="What is the capital of France?",
        options=["Berlin", "Madrid", "Paris", "Rome"],
        correct_answer=None  # None is allowed
    )
    assert question.correct_answer is None

# Test docx file reading
def test_read_docx():
    # Create a sample .docx file for testing
    doc = docx_instance()
    doc.add_paragraph("Q1. What is the capital of France?")
    doc.add_paragraph("Berlin")
    doc.add_paragraph("Madrid")
    doc.add_paragraph("Paris")
    doc.add_paragraph("Rome")
    doc.save("tests/data/sample.docx")
    
    result = read_docx("tests/data/sample.docx")
    assert isinstance(result, Document)

# Test paragraph parsing logic
def test_parse_para():
    doc = docx_instance()
    doc.add_paragraph("Q1. What is the capital of France?")
    doc.add_paragraph("Berlin")
    doc.add_paragraph("Madrid")
    doc.add_paragraph("Paris")
    doc.add_paragraph("Rome")
    parsed_questions = parse_para(doc)
    
    assert len(parsed_questions) == 1
    assert parsed_questions[0].question_text == "What is the capital of France?"
    assert parsed_questions[0].options == ["Berlin", "Madrid", "Paris", "Rome"]

# Test utility functions
def test_remove_prefix():
    assert remove_prefix("Q1. What is the capital of France?") == "What is the capital of France?"
    assert remove_prefix("1. What is the capital of France?") == "What is the capital of France?"

def test_find_marked_answer():
    doc = docx_instance()
    para = doc.add_paragraph("Berlin")
    para.runs[0].font.highlight_color = WD_COLOR.BLUE # Mark the answer
    
    answer = find_marked_answer(para, 1)
    assert answer == 0  # Marked answer should return index 0

# Test Excel writer functionality
def test_para_to_excel():
    questions = [
        Question(
            question_text="What is the capital of France?",
            options=["Berlin", "Madrid", "Paris", "Rome"],
            correct_answer=2
        )
    ]
    save_path = "tests/data/output.xlsx"
    para_to_excel(questions, save_path, columns=['Question', 'Option A', 'Option B', 'Option C', 'Option D', 'Correct Answer'])
    
    assert os.path.exists(save_path)
    
    df = pd.read_excel(save_path)
    assert len(df) == 1
    assert df.iloc[0]['Question'] == "What is the capital of France?"
    assert df.iloc[0]['Correct Answer'] == 2
